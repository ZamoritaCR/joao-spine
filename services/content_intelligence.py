"""JOAO Universal Content Intelligence — YouTube, PDF, Excel, DOCX, web links."""

from __future__ import annotations

import logging
import os
import re
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

import anthropic

logger = logging.getLogger(__name__)

_PROJECT_ROOT = Path(__file__).resolve().parent.parent
_LINKS_DIR = _PROJECT_ROOT / "links"
_OUTPUTS_DIR = _PROJECT_ROOT / "outputs"
_SESSION_LOG = _PROJECT_ROOT / "JOAO_SESSION_LOG.md"

_LINKS_DIR.mkdir(parents=True, exist_ok=True)
_OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)


def _ts() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")


def _claude(prompt: str, content: str, max_tokens: int = 4096) -> str:
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    msg = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=max_tokens,
        messages=[{"role": "user", "content": f"{prompt}\n\n{content}"}],
    )
    return msg.content[0].text


def _append_session_log(header: str, body: str) -> None:
    ts = _ts()
    entry = f"\n## {header} — {ts}\n\n{body}\n"
    try:
        with _SESSION_LOG.open("a", encoding="utf-8") as f:
            f.write(entry)
    except Exception as e:
        logger.warning("Failed to append session log: %s", e)


def _update_outputs_index() -> None:
    """Regenerate the outputs index.html from all files in _OUTPUTS_DIR."""
    files = sorted(_OUTPUTS_DIR.glob("*.html"), key=lambda p: p.stat().st_mtime, reverse=True)
    rows = ""
    for f in files:
        if f.name == "index.html":
            continue
        mtime = datetime.fromtimestamp(f.stat().st_mtime, tz=timezone.utc).strftime(
            "%Y-%m-%d %H:%M UTC"
        )
        name = f.name
        ftype = "Unknown"
        if name.startswith("pdf_"):
            ftype = "PDF"
        elif name.startswith("excel_") or name.startswith("csv_"):
            ftype = "Excel/CSV"
        elif name.startswith("docx_"):
            ftype = "Word"
        elif name.startswith("youtube_"):
            ftype = "YouTube"
        rows += (
            f'<tr><td><a href="{name}">{name}</a></td>'
            f"<td>{ftype}</td><td>{mtime}</td></tr>\n"
        )

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>JOAO Intelligence Reports</title>
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: 'Segoe UI', system-ui, sans-serif; background: #0d0d0f; color: #e2e8f0; padding: 40px; }}
  h1 {{ font-size: 1.8rem; font-weight: 700; margin-bottom: 8px; color: #fff; }}
  p.sub {{ color: #64748b; font-size: 0.9rem; margin-bottom: 32px; }}
  table {{ width: 100%; border-collapse: collapse; background: #161618; border-radius: 10px; overflow: hidden; }}
  th {{ background: #1e1e24; color: #94a3b8; font-size: 0.75rem; text-transform: uppercase;
        letter-spacing: 0.08em; padding: 12px 16px; text-align: left; }}
  td {{ padding: 12px 16px; border-top: 1px solid #1e1e24; font-size: 0.9rem; }}
  td a {{ color: #818cf8; text-decoration: none; }}
  td a:hover {{ text-decoration: underline; }}
  tr:hover td {{ background: #1a1a20; }}
</style>
</head>
<body>
<h1>JOAO Intelligence Reports</h1>
<p class="sub">Auto-generated. Refreshed on each new report. {len(files)} report(s) total.</p>
<table>
<thead><tr><th>File</th><th>Type</th><th>Generated</th></tr></thead>
<tbody>
{rows if rows else '<tr><td colspan="3" style="color:#475569;text-align:center;padding:40px">No reports yet.</td></tr>'}
</tbody>
</table>
</body>
</html>"""

    index_path = _OUTPUTS_DIR / "index.html"
    index_path.write_text(html, encoding="utf-8")


# ── YouTube ──────────────────────────────────────────────────────────────────

_YT_PATTERNS = [
    re.compile(r"(?:v=|youtu\.be/|/embed/|/v/|/shorts/)([A-Za-z0-9_-]{11})"),
]


def _extract_video_id(url: str) -> str | None:
    for pat in _YT_PATTERNS:
        m = pat.search(url)
        if m:
            return m.group(1)
    return None


def _fetch_youtube_transcript(video_id: str) -> str:
    from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound

    try:
        segments = YouTubeTranscriptApi.get_transcript(video_id)
        return " ".join(s["text"] for s in segments)
    except (TranscriptsDisabled, NoTranscriptFound):
        return ""
    except Exception as e:
        logger.warning("youtube-transcript-api error for %s: %s", video_id, e)
        return ""


def _fetch_youtube_whisper(url: str, video_id: str) -> str:
    """Download audio with yt-dlp then transcribe with whisper."""
    audio_path = _LINKS_DIR / f"yt_audio_{video_id}.mp3"
    try:
        subprocess.run(
            [
                sys.executable, "-m", "yt_dlp",
                "--extract-audio", "--audio-format", "mp3",
                "--output", str(audio_path),
                "--quiet",
                url,
            ],
            check=True,
            timeout=300,
        )
    except Exception as e:
        raise RuntimeError(f"yt-dlp download failed: {e}") from e

    try:
        import whisper

        model = whisper.load_model("base")
        result = model.transcribe(str(audio_path))
        return result["text"]
    except Exception as e:
        raise RuntimeError(f"Whisper transcription failed: {e}") from e
    finally:
        if audio_path.exists():
            audio_path.unlink()


async def handle_youtube(url: str) -> dict:
    video_id = _extract_video_id(url)
    if not video_id:
        raise ValueError(f"Could not extract video ID from URL: {url}")

    ts = _ts()

    # Step 1: get transcript
    transcript = _fetch_youtube_transcript(video_id)
    if not transcript:
        logger.info("No auto-transcript for %s — falling back to whisper", video_id)
        transcript = _fetch_youtube_whisper(url, video_id)

    if not transcript:
        raise RuntimeError("Could not obtain transcript for this video.")

    # Step 2: write raw transcript
    raw_path = _LINKS_DIR / f"youtube_{video_id}_{ts}.txt"
    raw_path.write_text(transcript, encoding="utf-8")

    # Step 3: Claude analysis
    prompt = (
        "You are JOAO's learning engine. Extract every insight, concept, framework, and "
        "actionable idea from this content. Then identify exactly how each insight applies "
        "to these active projects: dopamine.watch, dopamine.chat, TAOP Connect, Dr. Data, JOAO. "
        "Be specific. No fluff. No emojis."
    )
    analysis = _claude(prompt, transcript[:50000], max_tokens=4096)

    # Step 4: append to SESSION_LOG
    _append_session_log(f"LEARNED: {url}", analysis)

    return {
        "type": "youtube",
        "video_id": video_id,
        "transcript_path": str(raw_path),
        "analysis": analysis,
    }


# ── Universal Web Link ────────────────────────────────────────────────────────

async def handle_web_link(url: str) -> dict:
    text = ""

    # Try requests + BeautifulSoup first
    try:
        import requests
        from bs4 import BeautifulSoup

        r = requests.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
    except Exception as e:
        logger.warning("BeautifulSoup fetch failed for %s: %s", url, e)

    # Fallback to trafilatura
    if not text or len(text) < 200:
        try:
            import trafilatura

            downloaded = trafilatura.fetch_url(url)
            if downloaded:
                extracted = trafilatura.extract(downloaded)
                if extracted:
                    text = extracted
        except Exception as e:
            logger.warning("trafilatura fallback failed for %s: %s", url, e)

    if not text:
        raise RuntimeError(f"Could not extract content from {url}")

    ts = _ts()
    slug = re.sub(r"[^a-zA-Z0-9]", "_", url)[:60]
    raw_path = _LINKS_DIR / f"web_{slug}_{ts}.txt"
    raw_path.write_text(text, encoding="utf-8")

    prompt = (
        "You are JOAO's learning engine. Read this web page completely. "
        "Extract every insight, concept, framework, and actionable idea. "
        "Then identify exactly how each insight applies to these active projects: "
        "dopamine.watch, dopamine.chat, TAOP Connect, Dr. Data, JOAO. "
        "Be specific. No fluff. No emojis."
    )
    analysis = _claude(prompt, text[:50000], max_tokens=4096)

    _append_session_log(f"LEARNED: {url}", analysis)

    return {
        "type": "web_link",
        "url": url,
        "raw_path": str(raw_path),
        "analysis": analysis,
    }


# ── PDF ───────────────────────────────────────────────────────────────────────

async def handle_pdf(filename: str, data: bytes) -> dict:
    import pdfplumber
    import io

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages_text = []
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                pages_text.append(t)
    raw_text = "\n\n".join(pages_text)

    ts = _ts()
    base = Path(filename).stem
    raw_path = _LINKS_DIR / f"pdf_{base}_{ts}.txt"
    raw_path.write_text(raw_text, encoding="utf-8")

    prompt = (
        "Read this document completely. Then do two things:\n"
        "First, produce the most visually stunning, information-dense HTML report summarizing "
        "everything — use clean modern CSS, structured sections, tables where appropriate, zero emojis.\n"
        "Second, identify every insight that applies to TAOP or JOAO projects and list them explicitly.\n"
        "Output ONLY the complete HTML document. No preamble, no markdown fences."
    )
    html_output = _claude(prompt, raw_text[:60000], max_tokens=8192)

    html_path = _OUTPUTS_DIR / f"pdf_{base}_{ts}.html"
    html_path.write_text(html_output, encoding="utf-8")
    _update_outputs_index()

    # Learning summary for session log
    summary_prompt = (
        "Summarize the 5 most important insights from this document that apply to "
        "TAOP or JOAO projects. Plain text, numbered list, concise."
    )
    learning_summary = _claude(summary_prompt, raw_text[:30000], max_tokens=1024)
    _append_session_log(f"LEARNED: {filename}", learning_summary)

    return {
        "type": "pdf",
        "filename": filename,
        "raw_path": str(raw_path),
        "html_path": str(html_path),
        "html_url": f"/joao/outputs/{html_path.name}",
        "learning_summary": learning_summary,
    }


# ── Excel / CSV ───────────────────────────────────────────────────────────────

async def handle_spreadsheet(filename: str, data: bytes) -> dict:
    import pandas as pd
    import io

    ext = Path(filename).suffix.lower()
    base = Path(filename).stem
    ts = _ts()

    if ext == ".csv":
        df_dict = {"Sheet1": pd.read_csv(io.BytesIO(data))}
    else:
        xl = pd.ExcelFile(io.BytesIO(data))
        df_dict = {sheet: xl.parse(sheet) for sheet in xl.sheet_names}

    # Build analysis profile
    profile_parts = []
    for sheet_name, df in df_dict.items():
        part = f"SHEET: {sheet_name}\n"
        part += f"Rows: {len(df)} | Columns: {len(df.columns)}\n"
        part += f"Headers: {list(df.columns)}\n"
        null_counts = df.isnull().sum()
        part += f"Null counts: {null_counts.to_dict()}\n"
        dtypes = {col: str(dtype) for col, dtype in df.dtypes.items()}
        part += f"Data types: {dtypes}\n"
        part += f"Sample rows (first 10):\n{df.head(10).to_string()}\n"
        numeric_cols = df.select_dtypes(include="number")
        if not numeric_cols.empty:
            stats = numeric_cols.agg(["min", "max", "mean"]).to_string()
            part += f"Numeric stats:\n{stats}\n"
        profile_parts.append(part)

    profile_text = "\n\n".join(profile_parts)

    raw_path = _LINKS_DIR / f"{'csv' if ext == '.csv' else 'excel'}_{base}_{ts}.txt"
    raw_path.write_text(profile_text, encoding="utf-8")

    prompt = (
        "You are Dr. Data. You have the brain of a world-class data analyst and business "
        "intelligence expert. Analyze this dataset completely. Identify patterns, anomalies, "
        "business insights, data quality issues, and recommended visualizations. Then produce "
        "a full HTML report — stunning design, clean CSS, interactive-looking tables, summary "
        "stats, insight cards, recommendations section. No emojis. This should look like a "
        "Fortune 500 BI dashboard report. Output ONLY the complete HTML document."
    )
    html_output = _claude(prompt, profile_text[:50000], max_tokens=8192)

    prefix = "csv_" if ext == ".csv" else "excel_"
    html_path = _OUTPUTS_DIR / f"{prefix}{base}_{ts}.html"
    html_path.write_text(html_output, encoding="utf-8")
    _update_outputs_index()

    insights_prompt = (
        "Based on this dataset profile, list the 5 most important business insights or "
        "anomalies found. Plain text, numbered, concise."
    )
    insights = _claude(insights_prompt, profile_text[:20000], max_tokens=1024)
    _append_session_log(f"DR. DATA: {filename}", insights)

    return {
        "type": "spreadsheet",
        "filename": filename,
        "raw_path": str(raw_path),
        "html_path": str(html_path),
        "html_url": f"/joao/outputs/{html_path.name}",
        "insights": insights,
    }


# ── DOCX ──────────────────────────────────────────────────────────────────────

async def handle_docx(filename: str, data: bytes) -> dict:
    from docx import Document
    import io

    doc = Document(io.BytesIO(data))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else ""
            if "Heading" in style:
                parts.append(f"\n### {para.text.strip()}\n")
            else:
                parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    raw_text = "\n".join(parts)

    ts = _ts()
    base = Path(filename).stem
    raw_path = _LINKS_DIR / f"docx_{base}_{ts}.txt"
    raw_path.write_text(raw_text, encoding="utf-8")

    prompt = (
        "Read this document. Produce the most amazing HTML version of it — preserve structure, "
        "headings, tables, make it visually stunning with modern CSS. Then extract all insights "
        "applicable to TAOP or JOAO. Output ONLY the complete HTML document. No preamble."
    )
    html_output = _claude(prompt, raw_text[:60000], max_tokens=8192)

    html_path = _OUTPUTS_DIR / f"docx_{base}_{ts}.html"
    html_path.write_text(html_output, encoding="utf-8")
    _update_outputs_index()

    insights_prompt = (
        "Summarize the key insights from this Word document that apply to TAOP or JOAO. "
        "Plain text, numbered list, concise."
    )
    insights = _claude(insights_prompt, raw_text[:30000], max_tokens=1024)
    _append_session_log(f"LEARNED: {filename}", insights)

    return {
        "type": "docx",
        "filename": filename,
        "raw_path": str(raw_path),
        "html_path": str(html_path),
        "html_url": f"/joao/outputs/{html_path.name}",
        "insights": insights,
    }
